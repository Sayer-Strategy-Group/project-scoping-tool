from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('The Milestone Group — HubSpot CRM Proposal', level=1)

# ============================================================
# A NOTE ON THIS PROPOSAL
# ============================================================
doc.add_heading('A Note on This Proposal', level=3)

doc.add_paragraph(
    'Proposals are inherently a starting point—not a final blueprint. This document reflects our '
    'understanding of your challenges, a proposed solution, and the resources required to execute successfully.'
)
doc.add_paragraph(
    'We recognize that you, as the client, live with these challenges daily, and there may be details '
    'we have yet to uncover. Our goal is to use this as a conversation starter—to refine, align, and '
    'ensure we are solving the right problems in the right order.'
)
doc.add_paragraph(
    'If any aspect of the scope, solution, or pricing feels misaligned, let\'s use that as a basis '
    'for further discussion.'
)

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=2)

doc.add_paragraph(
    'The Milestone Group is a national private equity real estate investment management firm with '
    'industry-leading returns since inception in 2003 and over $7 billion in real estate investments. '
    'With approximately 45 team members across offices in Dallas, Boca Raton, Atlanta, and a fourth location, '
    'Milestone manages apartment acquisitions across major U.S. metropolitan markets on behalf of leading '
    'institutional investors, pension funds, and sovereign wealth funds.'
)
doc.add_paragraph(
    'This proposal outlines a phased HubSpot CRM implementation designed to centralize contact and relationship '
    'management, capture deal activity, and replace fragmented tracking across Excel, OneNote, and individual '
    'email accounts. The engagement is structured as a three-phase rollout: Phase 1 delivers contact management '
    'and activity capture (weeks 1–4), Phase 2 introduces the deal pipeline (weeks 5–6), and Phase 3 adds '
    'integrations and advanced features (weeks 7–8+).'
)
doc.add_paragraph(
    'This phased approach is strongly recommended based on Milestone\'s stated preference for starting simple '
    'and expanding. Team members range from 25 to 65 years old with varying technology comfort levels—a gradual '
    'rollout allows less technical users to build confidence with basic contact management before adding deal '
    'pipeline complexity. The previous DealCloud abandonment after two months underscores the adoption risk of '
    'an all-at-once approach. Phase 1 delivers immediate value with minimal behavior change, while Phases 2 '
    'and 3 build on established habits.'
)

# ============================================================
# 2. OBJECTIVES
# ============================================================
doc.add_heading('2. Objectives', level=2)

objectives = [
    'Centralize contact and company data across all four offices into a single HubSpot CRM, replacing '
    'fragmented Excel, OneNote, and email-based tracking.',
    'Configure relationship management for key personas—brokers, evangelists, investment partners, and '
    'lenders—with parent-child company associations for PE firm-to-portfolio company structures.',
    'Capture and log all business development activity (calls, meetings, emails) automatically through '
    'Google Workspace and Fireflies integrations.',
    'Build a structured deal pipeline with stage gates and qualification criteria to track opportunities '
    'from initial contact through close.',
    'Deliver reporting dashboards that provide executive visibility into BD activity, contact engagement, '
    'and pipeline health.',
    'Migrate and deduplicate 2,000–5,000 existing contacts from multiple sources into a clean, reliable '
    'CRM foundation.',
    'Equip the team with training, documentation, and mobile access to drive adoption across all offices '
    'and age groups.',
]
for obj in objectives:
    doc.add_paragraph(obj)

# ============================================================
# 3. SCOPE OF WORK
# ============================================================
doc.add_heading('3. Scope of Work', level=2)
doc.add_paragraph('The following workstreams and tasks will be delivered during this engagement:')

# -- Helper function for workstreams --
def add_workstream(title, customer_story, approaches, assumptions):
    doc.add_heading(title, level=3)
    
    doc.add_paragraph('Customer Story')
    doc.add_paragraph(customer_story)
    
    doc.add_paragraph('Recommended Approach')
    for a in approaches:
        doc.add_paragraph(a)
    
    doc.add_paragraph('Assumptions')
    for a in assumptions:
        doc.add_paragraph(a)

# WORKSTREAM 1
add_workstream(
    'Workstream 1: CRM Architecture & Setup',
    'The Milestone Group wants a professionally configured HubSpot CRM environment that reflects '
    'their brand and supports their multi-office team structure with appropriate access controls.',
    [
        'Configure HubSpot portal with Milestone branding and company settings',
        'Set up user roles and permission sets for up to 10 users across 4 offices',
        'Configure team structure aligned to office locations and functional roles',
        'Establish naming conventions and organizational standards for long-term CRM hygiene',
        'Configure security settings and data access controls appropriate for a PE firm',
    ],
    [
        'Assumes HubSpot Sales Hub Starter or Smart CRM tier; final tier confirmed after scoping call',
        'Up to 10 licensed CRM user seats',
        'Portal configuration includes branding, default settings, and user provisioning',
    ]
)

# WORKSTREAM 2
add_workstream(
    'Workstream 2: Contact & Company Management',
    'The Milestone Group wants to track relationships across brokers, evangelists, investment partners, '
    'lenders, and portfolio companies in a structured system that reflects the complexity of PE real estate '
    'relationship networks.',
    [
        'Configure custom contact and company properties tailored to real estate PE workflows',
        'Set up lifecycle stages for prospect-to-active relationship progression',
        'Create association labels and parent-child company structures for PE firm-to-portfolio company relationships',
        'Define and configure persona categories: broker, evangelist, investment partner, lender',
        'Build segmentation lists for key contact categories and relationship types',
        'Establish deduplication strategy and detection rules for ongoing data quality',
    ],
    [
        'Includes up to 30 custom contact and company properties',
        'Parent-child associations configured for PE firm-to-portfolio company structures',
        'Persona tagging system for four primary relationship types',
        'Deduplication rules configured for automated detection',
    ]
)

# WORKSTREAM 3
add_workstream(
    'Workstream 3: Data Migration — Contact Import',
    'The Milestone Group wants to consolidate contacts currently scattered across Excel spreadsheets, '
    'OneNote, and individual email accounts into a single, clean CRM database.',
    [
        'Audit existing contact data across all sources for volume, quality, and completeness',
        'Develop field mapping between source data and HubSpot contact/company properties',
        'Cleanse, standardize, and deduplicate contact records in a staging environment',
        'Execute one test data load to validate mapping, formatting, and association accuracy',
        'Perform final production import with post-migration validation and cleanup',
    ],
    [
        'Assumes 2,000–5,000 contacts across 3–4 data sources',
        'Moderate deduplication effort expected given multiple overlapping sources',
        'One test load before final migration cutover',
        'Client to provide complete contact exports from Excel, OneNote, and email by week 2',
        'Does not include historical activity or deal data migration',
    ]
)

# WORKSTREAM 4
add_workstream(
    'Workstream 4: Email & Communication Setup',
    'The Milestone Group wants seamless integration between their Google Workspace environment and '
    'HubSpot so that emails, meetings, and calendar activity are automatically captured in the CRM.',
    [
        'Configure Google Workspace integration for Gmail and Google Calendar per user',
        'Set up connected inbox configuration for each CRM user',
        'Configure email tracking settings with appropriate opt-in preferences',
        'Establish email domain exclusions to prevent logging of internal or irrelevant communications',
        'Set up calendar sync for automatic meeting capture and logging',
    ],
    [
        'Assumes Google Workspace as primary email and calendar platform',
        'Per-user inbox connection and email tracking configuration',
        'Includes email domain exclusion rules for internal communications',
        'Calendar and contact sync configured across all active CRM users',
    ]
)

# WORKSTREAM 5
add_workstream(
    'Workstream 5: Automations & Workflows',
    'The Milestone Group wants to reduce manual data entry and ensure consistent follow-up by automating '
    'activity capture, task reminders, and contact creation rules.',
    [
        'Configure automated activity capture rules for emails and meetings',
        'Build task reminder triggers for follow-up actions and overdue activities',
        'Create follow-up workflows based on meeting outcomes and contact engagement',
        'Set up automatic contact creation rules with appropriate filters',
        'Configure email domain exclusions for auto-contact creation',
    ],
    [
        'Includes up to 8 automated workflows',
        'Covers activity capture automation, task reminders, and follow-up triggers',
        'Auto-contact creation rules include email domain exclusions',
        'Workflows will be documented for ongoing admin management',
    ]
)

# WORKSTREAM 6
add_workstream(
    'Workstream 6: Pipeline & Deal Management',
    'The Milestone Group wants a structured deal pipeline to track investment opportunities from initial '
    'qualification through close, with stage gates that enforce data discipline.',
    [
        'Build custom deal pipeline with stages aligned to Milestone\'s investment process',
        'Configure required fields per deal stage to enforce data completeness',
        'Set up qualification criteria and stage gate enforcement',
        'Create deal views and filters for team-based pipeline management',
        'Configure deal properties for investment-specific data points',
    ],
    [
        'Phase 2 deliverable (weeks 5–6); deal stages to be confirmed during scoping call',
        'Assumes 1 primary pipeline with 4–6 stages',
        'Includes required field enforcement per stage',
        'Stage progression rules configured for data quality',
    ]
)

# WORKSTREAM 7
add_workstream(
    'Workstream 7: Integration — Fireflies',
    'The Milestone Group wants meeting transcriptions from Fireflies to automatically log in HubSpot, '
    'reducing manual note-taking and ensuring institutional knowledge is captured in the CRM.',
    [
        'Configure Fireflies-to-HubSpot native integration',
        'Set up transcript mapping and activity logging rules',
        'Configure contact matching logic to associate transcriptions with correct CRM records',
        'Validate integration with sample meetings before go-live',
        'Document known limitations and troubleshooting procedures',
    ],
    [
        'Native Fireflies integration; no custom middleware required',
        'Includes transcript mapping, contact matching, and activity logging configuration',
        'Client must confirm Fireflies subscription with API access',
        'Phase 3 deliverable (weeks 7–8+)',
    ]
)

# WORKSTREAM 8
add_workstream(
    'Workstream 8: Integration — Google Workspace',
    'The Milestone Group wants calendar events, email communications, and contact data to sync '
    'bidirectionally between Google Workspace and HubSpot across all users.',
    [
        'Configure calendar sync automation for all CRM users',
        'Set up contact sync rules between Google Contacts and HubSpot',
        'Establish sync frequency and conflict resolution rules',
        'Validate sync accuracy and troubleshoot per-user configuration issues',
    ],
    [
        'Scoped separately from email setup to cover calendar and contact sync automation',
        'Assumes Google Workspace as the primary platform',
        'Phase 3 deliverable (weeks 7–8+)',
        'Includes per-user sync validation',
    ]
)

# WORKSTREAM 9
add_workstream(
    'Workstream 9: Reporting & Dashboards',
    'The Milestone Group wants executive-level visibility into business development activity, '
    'contact engagement, and deal pipeline health through purpose-built dashboards.',
    [
        'Build BD activity dashboard tracking calls, meetings, and emails by team member',
        'Create contact overview dashboard with engagement metrics and lifecycle tracking',
        'Configure deal pipeline dashboard with stage conversion and velocity metrics',
        'Set up up to 10 custom reports across all dashboards',
        'Configure report scheduling and email delivery for key stakeholders',
    ],
    [
        'Includes 2–3 dashboards: BD activity, contact overview, deal pipeline',
        'Up to 10 custom reports across all dashboards',
        'Uses standard HubSpot reporting capabilities',
        'Historical data available only from go-live date forward',
    ]
)

# WORKSTREAM 10
add_workstream(
    'Workstream 10: Training & Enablement',
    'The Milestone Group wants their team—ranging from highly technical to less tech-savvy members '
    'across four offices—to be confident and competent using HubSpot in their daily workflows.',
    [
        'Deliver 1–2 admin training sessions for designated power users covering settings, properties, '
        'workflows, and reporting',
        'Deliver 2–3 end-user training sessions tailored to varying technical comfort levels',
        'Cover daily CRM usage, mobile app setup, activity logging, and task management',
        'Provide hands-on guidance for each office\'s specific workflow patterns',
        'Record sessions for future onboarding of new team members',
    ],
    [
        '1–2 admin training sessions for 1–2 designated administrators',
        '2–3 end-user sessions designed for varying tech comfort levels (ages 25–65)',
        'Includes HubSpot mobile app setup and walkthrough',
        'Sessions delivered virtually via screen share',
    ]
)

# WORKSTREAM 11
add_workstream(
    'Workstream 11: UAT, Go-Live & Documentation',
    'The Milestone Group wants a controlled go-live process with testing, documentation, and post-launch '
    'support to ensure a smooth transition and sustained adoption.',
    [
        'Develop test scripts for user acceptance testing across all configured workstreams',
        'Conduct 2 UAT sessions with key stakeholders to validate configuration',
        'Execute bug fix window for issues identified during testing',
        'Create go-live cutover plan and checklist',
        'Deliver process documentation: 1 admin guide, 1 end-user guide, 1 mobile quick-start guide',
        'Provide 2-week hypercare support post-launch for issue resolution and adoption coaching',
    ],
    [
        'Includes test script creation, 2 UAT sessions, and bug fix window',
        'Go-live checklist and cutover plan included',
        'Deliverables: 1 admin guide, 1 end-user guide, 1 mobile quick-start guide',
        '2 weeks post-launch hypercare support included',
        'Ongoing managed services beyond hypercare are not included',
    ]
)

# ============================================================
# 4. DELIVERABLES & TIMELINE
# ============================================================
doc.add_heading('4. Deliverables & Timeline', level=2)

doc.add_paragraph(
    'This engagement will be completed over a 6–8 week period using a phased rollout approach. '
    'During that time, Sayer will:'
)

deliverables_list = [
    'Configure HubSpot CRM environment with Milestone branding, user roles, and team structure',
    'Migrate and deduplicate 2,000–5,000 contacts from Excel, OneNote, and email sources',
    'Set up Google Workspace integration for email, calendar, and activity capture',
    'Build automated workflows for activity logging, task reminders, and follow-up',
    'Configure deal pipeline with stage gates and qualification criteria',
    'Integrate Fireflies for automated meeting transcription logging',
    'Deliver reporting dashboards for BD activity, contact engagement, and pipeline visibility',
    'Train admin and end-user teams with role-appropriate sessions',
    'Provide documentation, UAT, go-live support, and 2-week hypercare',
]
for d in deliverables_list:
    doc.add_paragraph(d)

# Timeline table
doc.add_paragraph('')
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Phase / Week Range', 'Scope Highlights']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

timeline_data = [
    ('Phase 1: Weeks 1–4', 'CRM architecture & setup, contact/company management, data migration, '
     'email & communication setup, automations & workflows'),
    ('Phase 2: Weeks 5–6', 'Deal pipeline configuration, stage gates, qualification criteria, '
     'pipeline reporting'),
    ('Phase 3: Weeks 7–8+', 'Fireflies integration, Google Workspace sync, advanced reporting, '
     'dashboard refinement'),
    ('Go-Live: Weeks 6–8', 'UAT testing, training sessions, documentation delivery, go-live cutover, '
     'hypercare support'),
]
for i, (phase, highlights) in enumerate(timeline_data):
    table.rows[i+1].cells[0].text = phase
    table.rows[i+1].cells[1].text = highlights

# ============================================================
# 5. ENGAGEMENT MODEL & PRICING
# ============================================================
doc.add_heading('5. Engagement Model & Pricing', level=2)

# Pricing table
price_table = doc.add_table(rows=2, cols=3)
price_table.style = 'Table Grid'
price_table.alignment = WD_TABLE_ALIGNMENT.CENTER

price_headers = ['Component', 'Model', 'Fee']
for i, h in enumerate(price_headers):
    cell = price_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

price_table.rows[1].cells[0].text = 'Project Delivery'
price_table.rows[1].cells[1].text = 'Fixed Fee'
price_table.rows[1].cells[2].text = '$26,075'

doc.add_paragraph('')
doc.add_paragraph('Model: Fixed Fee')
doc.add_paragraph('Total Cost: $26,075')
doc.add_paragraph('Project Timeline: 6–8 weeks')

doc.add_paragraph('')
doc.add_paragraph('Payment Terms:')
doc.add_paragraph('50% ($13,037.50) invoiced at project start')
doc.add_paragraph('50% ($13,037.50) invoiced at final handoff (~Week 6–8)')
doc.add_paragraph('Net-15 terms on all invoices')

doc.add_paragraph('')
doc.add_paragraph(
    'Technology & Administrative Fee: To help offset the costs of our technology stack and internal '
    'systems that improve project speed, quality, and efficiency, a 5% Technology & Administrative '
    'fee is applied to each invoice.'
)

# ============================================================
# 6. PROJECT GOVERNANCE
# ============================================================
doc.add_heading('6. Project Governance', level=2)

doc.add_paragraph('Sayer Responsibilities')
sayer_resp = [
    'Lead configuration, QA, and implementation across all workstreams',
    'Conduct weekly check-ins and milestone reviews',
    'Provide documentation, training, and go-live support',
    'Manage project timeline and proactively communicate risks or blockers',
]
for r in sayer_resp:
    doc.add_paragraph(r)

doc.add_paragraph('')
doc.add_paragraph('The Milestone Group Responsibilities')
client_resp = [
    'Provide complete contact exports from Excel, OneNote, and email by week 2',
    'Designate 1 primary point of contact for decisions and approvals',
    'Ensure stakeholder availability for discovery, UAT, and training sessions',
    'Procure HubSpot licenses before implementation begins',
    'Procure or confirm Fireflies.ai subscription with API access',
    'Provide Google Workspace admin access for email/calendar integration setup',
    'Champion CRM adoption internally—leadership must set expectations for usage',
    'Complete UAT testing within agreed timeline',
    'Communicate deal stages and qualification criteria during scoping call',
]
for r in client_resp:
    doc.add_paragraph(r)

# ============================================================
# 7. ASSUMPTIONS & CONSTRAINTS
# ============================================================
doc.add_heading('7. Assumptions & Constraints', level=2)

assumptions = [
    'HubSpot Sales Hub Starter or Smart CRM tier (final tier confirmed after scoping call)',
    'Up to 10 licensed CRM users across 4 offices',
    'Google Workspace (Gmail + Google Calendar) as primary email/calendar platform',
    'Up to 30 custom contact/company properties',
    '1 deal pipeline with 4–6 stages (Phase 2)',
    'Up to 8 automated workflows',
    '2–3 reporting dashboards with up to 10 custom reports',
    '2,000–5,000 contacts to migrate from 3–4 sources',
    'Fireflies.ai as transcription tool (existing subscription)',
    '6–8 week implementation timeline with phased rollout',
    'All work will be conducted remotely unless otherwise agreed',
    'Any scope changes will be managed through a formal change request and may impact cost or timing',
]
for a in assumptions:
    doc.add_paragraph(a)

doc.add_paragraph('')
doc.add_paragraph('Out of Scope:')
out_of_scope = [
    'Salesforce, Zoho, or any CRM other than HubSpot',
    'ERP or accounting system integration',
    'Asana integration (recommended for future phase)',
    'Advanced deal flow management beyond basic pipeline (e.g., deal scoring, weighted pipeline)',
    'Custom API development or middleware (n8n/Zapier)',
    'Marketing automation (email campaigns, sequences, landing pages)',
    'DocuSign or e-signature integration',
    'HubSpot Service Hub or Operations Hub features',
    'Ongoing managed services beyond 2-week hypercare',
    'Hardware procurement or IT infrastructure changes',
]
for o in out_of_scope:
    doc.add_paragraph(o)

# ============================================================
# 8. APPROVAL & NEXT STEPS
# ============================================================
doc.add_heading('8. Approval & Next Steps', level=2)

doc.add_paragraph(
    'Please confirm alignment on scope and structure so we can move forward with scheduling a '
    'formal kickoff and assigning your Sayer team.'
)
doc.add_paragraph(
    'If any scope areas, phasing, or pricing need to be adjusted to fit your needs, we welcome '
    'the conversation.'
)
doc.add_paragraph('')
doc.add_paragraph('To move forward with this engagement:')
doc.add_paragraph('Provide written approval of this proposal via email')
doc.add_paragraph('Schedule project kickoff within 1–2 weeks of written approval')
doc.add_paragraph('')
doc.add_paragraph(
    'Upon written approval, Sayer will issue a Master Services Agreement (MSA) via DocuSign, '
    'with this Scope of Work included. Once the MSA is executed, Sayer will finalize the '
    'implementation plan and begin onboarding.'
)

# Save
output_path = '/Users/harbuckconsulting/Project Scoping Tool/The Milestone Group HubSpot CRM Proposal.docx'
doc.save(output_path)
print(f'Proposal saved to: {output_path}')
