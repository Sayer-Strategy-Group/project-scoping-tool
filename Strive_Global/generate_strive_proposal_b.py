"""Generate STRIVE Global -- Full CRM Overhaul Proposal (Option B)"""
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'


def add_workstream(title, customer_story, approaches, assumptions):
    doc.add_heading(title, level=3)
    doc.add_paragraph('Customer Story')
    doc.add_paragraph(customer_story)
    doc.add_paragraph('Recommended Approach')
    for a in approaches:
        doc.add_paragraph(a)
    doc.add_paragraph('Assumptions')
    for a in assumptions:
        doc.add_paragraph(a)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph('')


# ============================================================
# TITLE
# ============================================================
doc.add_heading('STRIVE Global -- Full HubSpot CRM Overhaul Proposal', level=1)
doc.add_paragraph('Option B: CRM Foundation + CPQ, Enrichment, Marketing & Integration Prep')

# ============================================================
# A NOTE ON THIS PROPOSAL
# ============================================================
doc.add_heading('A Note on This Proposal', level=3)
doc.add_paragraph(
    'Proposals are inherently a starting point -- not a final blueprint. This document reflects our '
    'understanding of your challenges, a proposed solution, and the resources required to execute successfully.'
)
doc.add_paragraph(
    'We recognize that you, as the client, live with these challenges daily, and there may be details '
    'we have yet to uncover. Our goal is to use this as a conversation starter -- to refine, align, and '
    'ensure we are solving the right problems in the right order.'
)
doc.add_paragraph(
    'If any aspect of the scope, solution, or pricing feels misaligned, let\'s use that as a basis '
    'for further discussion.'
)

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=2)

doc.add_paragraph(
    'STRIVE Global is a PE-backed (Essex Bay Capital) benefits technology company headquartered in Denver, CO '
    'with approximately 51 employees. STRIVE offers white-labeled employee benefits apps, AI-powered benefits '
    'assistants, scalable microsites, and a rewards platform (Divi) distributed through a broker channel model. '
    'Sales motions range from $6K client-level deals to $400K+ enterprise broker packs.'
)
doc.add_paragraph(
    'This proposal outlines a comprehensive CRM overhaul delivered in two phases. Phase 1 (weeks 1-6) '
    'addresses the CRM foundation: pipeline restructuring, contact classification, deal health scoring, and '
    'automated board dashboards. Phase 2 (weeks 7-14) builds on the clean foundation with CPQ implementation, '
    'AI-powered data enrichment at scale, email sequencing, marketing automation, commission visibility, and '
    'integration preparation for NetSuite and Bill.com.'
)
doc.add_paragraph(
    'The phased approach is strongly recommended. STRIVE\'s data quality issues make it risky to build CPQ on '
    'unclassified contacts and corrupted pipeline data. Phase 1 creates the clean foundation that Phase 2 '
    'depends on. This also aligns with Jodi\'s stated goal of being cost-effective this year while building '
    'the foundation for future scaling -- Phase 1 can stand alone if Phase 2 is deferred.'
)
doc.add_paragraph(
    'The total investment pays for itself within 4-7 months through quantifiable time savings and tool cost '
    'reduction, with additional upside from pipeline visibility, cross-sell automation, and board confidence.'
)

# ============================================================
# 2. OBJECTIVES
# ============================================================
doc.add_heading('2. Objectives', level=2)

objectives = [
    'Separate pre-pipeline deals from active pipeline to eliminate metric distortion and enable accurate forecasting.',
    'Classify all contact records by type using AI enrichment, unlocking segmented marketing and targeted outreach.',
    'Implement deal health scoring to replace gut-feel forecasting with objective pipeline reviews.',
    'Deliver automated dashboards for partner, seller, and CEO/board reporting -- eliminating 8-10 hours/month of manual assembly.',
    'Implement HubSpot CPQ with a complete product library, quote templates, e-signatures, and automated reminders '
    '-- replacing manual quoting and capturing deals currently invisible in the pipeline.',
    'Build contract management workflows to replace the manual DocuSign process through the COO.',
    'Deploy AI-powered data enrichment at scale using HubSpot Breeze credits to auto-classify and enrich the full database.',
    'Establish parent-child company structures for broker pack tracking, enabling microsite-to-app expansion prospecting.',
    'Launch email sequencing for broker prospecting and client expansion outreach.',
    'Build marketing automation foundation: newsletter templates and segmented campaigns by contact type.',
    'Prepare data infrastructure for future NetSuite and Bill.com integration (prep only -- integration build separate).',
    'Audit and optimize HubSpot seat/license allocation to reduce unnecessary subscription costs.',
    'Implement commission visibility and validation workflow tied to closed-won deals.',
    'Standardize close won/lost reasons, lead source tracking, and data entry requirements across the team.',
]
for obj in objectives:
    doc.add_paragraph(obj)

# ============================================================
# 3. SCOPE OF WORK
# ============================================================
doc.add_heading('3. Scope of Work', level=2)
doc.add_paragraph('The following workstreams are organized by phase. Phase 1 delivers the CRM foundation; '
                   'Phase 2 builds CPQ, enrichment, marketing, and integration preparation on that foundation.')

# -- PHASE 1 --
doc.add_heading('Phase 1: CRM Foundation (Weeks 1-6)', level=2)

add_workstream(
    'Workstream 1: CRM Architecture & Configuration Review',
    'STRIVE\'s HubSpot instance needs auditing and optimization to support the team\'s actual workflows.',
    [
        'Audit current portal configuration, user roles, and permission sets',
        'Reconfigure team structure for partner vs. client sellers',
        'Review seat allocation across all hubs',
        'Establish naming conventions and organizational standards',
    ],
    ['Existing Enterprise instance. Up to 23 licensed users.']
)

add_workstream(
    'Workstream 2: Pipeline Restructuring',
    'Pre-pipeline deals mixed with active deals distort metrics. Partner and client deals need segmentation.',
    [
        'Separate pre-pipeline into its own pipeline or board view',
        'Configure partner-level and client-level deal categories',
        'Map stage probabilities (5/20/40/50/60/80/90/100)',
        'Add Contracting stage. Configure required fields per stage.',
        'Automate pre-pipeline to active pipeline transitions with deal age reset',
    ],
    ['Same stages for both categories. Product segmentation via product set fields.']
)

add_workstream(
    'Workstream 3: Contact & Company Classification + Cleanup',
    '10 years of uncoded contacts. Zero classification. Cannot segment for outreach or reporting.',
    [
        'Define contact type taxonomy (client, prospect, broker, partner)',
        'Run AI enrichment pilot on 500 records',
        'Build classification rules. Execute full database classification.',
        'Normalize, deduplicate, fix associations. Establish duplicate protection.',
    ],
    ['+30-50% data quality premium. Up to 30 custom properties. HubSpot Breeze credits.']
)

add_workstream(
    'Workstream 4: Deal Health Scoring & Forecasting',
    'No algorithmic forecasting today. Gut feel only.',
    [
        'Create deal score: stage + time-in-stage + activity recency',
        'Classify deals as Healthy, At Risk, or Stale',
        'Replace free-form close won/lost with categorized dropdowns',
    ],
    ['Scoring formula defined collaboratively. Industry benchmarks as starting point.']
)

add_workstream(
    'Workstream 5: Reporting & Dashboards',
    'Board reports take 1+ day. Andrea needs clonable partner dashboards. Jodi needs CEO-level visibility.',
    [
        'Build partner, seller, CEO/board, and pipeline health dashboards',
        'Configure up to 12 custom reports',
        'Set up report scheduling and export',
    ],
    ['3-4 dashboards. All Andrea\'s requested metrics included.']
)

add_workstream(
    'Workstream 6: Automations & Workflows',
    'Almost no automation beyond broker trial form. Outlook logging defect creating noise.',
    [
        'Fix Outlook email logging issue',
        'Lead source tracking. Form-to-CRM automation.',
        'Stage-change notifications. Duplicate protection.',
    ],
    ['Up to 8 workflows. Outlook fix may require HubSpot support.']
)

add_workstream(
    'Workstream 7: Email & Communication Fix',
    'Outlook emails log to every deal a contact is associated with.',
    [
        'Audit integration. Configure association rules.',
        'Domain exclusions. Validate with sample scenarios.',
    ],
    ['Root cause diagnosis week 1.']
)

add_workstream(
    'Workstream 8-10: Training (Admin, End Users, Executive)',
    'Admin, sales team, and Jodi each need role-appropriate training on the reconfigured CRM.',
    [
        '1-2 admin sessions: pipeline management, reporting, workflows, data hygiene',
        '2-3 end-user sessions: daily CRM usage, deal entry, activity logging, dashboards, mobile',
        '1 executive session: CEO dashboard, forecast consumption, board report scheduling',
    ],
    ['All sessions virtual. Recorded for future onboarding.']
)

add_workstream(
    'Workstream 11: UAT, Go-Live & Documentation (Phase 1)',
    'Controlled go-live with testing and 2-week hypercare.',
    [
        'Test scripts. 2 UAT sessions. Bug fix window.',
        'Go-live cutover plan. Admin guide + user guide.',
        '2-week hypercare post-launch.',
    ],
    ['Phase 1 go-live. Phase 2 has separate UAT cycle.']
)

# -- PHASE 2 --
doc.add_heading('Phase 2: CPQ, Enrichment & Integrations (Weeks 7-14)', level=2)

add_workstream(
    'Workstream 13: CPQ Implementation',
    'STRIVE manually creates every quote. Pricing is inconsistent. Deals are quoted outside the CRM and never '
    'enter the pipeline, causing systematic under-reporting. Jodi sees CPQ as critical to de-risking contracts '
    'and capturing accurate pipeline data.',
    [
        'Build complete product library (apps, microsites, AI assistant, broker packs, Divi/rewards, bundles)',
        'Configure pricing models: license fees, revenue share, flat per-head, bundle discounts',
        'Design quote templates branded to STRIVE',
        'Configure e-signature capability within HubSpot',
        'Set up automated quote reminders for unsigned quotes',
        'Build quote-to-deal automation (auto-update deal amount and products on quote acceptance)',
    ],
    [
        'Requires complete product/pricing inventory from Andrea before build begins',
        'Multiple pricing models supported: license + rev share, flat per-head, bundles/discounts',
        'Uses HubSpot Enterprise CPQ capabilities (included in current license)',
        '3-5 recent quotes across deal types needed as reference',
    ]
)

add_workstream(
    'Workstream 14: Contract Management Workflows',
    'Contracts currently go through COO via DocuSign with no CRM tracking. Status is invisible in HubSpot.',
    [
        'Build quote-to-contract automation workflow',
        'Configure contract status tracking properties on deals',
        'Set up renewal and expiry notification workflows',
        'Establish contracting stage requirements (auto-advance on signature)',
    ],
    ['Replaces manual DocuSign process. Contract templates managed in HubSpot.']
)

add_workstream(
    'Workstream 15: Data Enrichment at Scale',
    'Phase 1 pilots AI enrichment on 500 records. Phase 2 scales it to the full database and sets up ongoing '
    'auto-enrichment for new contacts.',
    [
        'Scale HubSpot Breeze AI enrichment to full contact database',
        'Configure auto-enrichment rules for new inbound contacts',
        'QA and validate enrichment results across contact types',
        'Document enrichment accuracy and coverage metrics',
    ],
    ['Uses HubSpot Breeze credits (5-10K). Extends Phase 1 pilot to full database.']
)

add_workstream(
    'Workstream 16: Data Normalization & Deduplication (Extended)',
    'Phase 1 handles basic cleanup. Phase 2 performs comprehensive database normalization.',
    [
        'Full database normalization (capitalization, phone formats, addresses)',
        'Advanced deduplication with configurable matching rules',
        'Merge strategy for duplicate records (preserve most complete record)',
        'Ongoing duplicate protection rules for sustained data quality',
    ],
    ['Extends Phase 1 cleanup. Full +50% data quality premium applied.']
)

add_workstream(
    'Workstream 17: Parent-Child Structure (Broker Packs)',
    'The Howden 250-pack deal needs tracking. Each microsite client should be a child of the broker parent, '
    'enabling app upsell prospecting.',
    [
        'Configure company-to-company association labels (parent broker, child client)',
        'Build per-broker intake forms in HubSpot CMS',
        'Automate child contact creation from form submissions',
        'Generate segmented lists by parent broker for expansion prospecting',
    ],
    ['Depends on microsite form link from Andrea. Enables crawl-walk-run product expansion tracking.']
)

add_workstream(
    'Workstream 18: Email Sequencing',
    'No email sequencing tool today. Reps do all outreach manually.',
    [
        'Build broker prospecting sequence (warm outreach to new broker contacts)',
        'Build microsite-to-app expansion sequence (upsell existing microsite clients)',
        'Configure up to 4 total sequences with enrollment triggers',
        'Set up sequence performance reporting',
    ],
    ['Uses HubSpot Sales Hub Enterprise sequences (included in current license). Up to 4 sequences.']
)

add_workstream(
    'Workstream 19: Marketing Automation Foundation',
    'STRIVE pays for Marketing Hub but runs zero campaigns. Contact classification (Phase 1) enables segmented outreach.',
    [
        'Design newsletter email template branded to STRIVE',
        'Build 2-3 segmented campaign workflows (by contact type: client, broker, prospect)',
        'Configure unsubscribe management and CAN-SPAM compliance',
        'Set up campaign performance reporting',
    ],
    ['Leverages Marketing Hub (already licensed). Depends on Phase 1 contact classification.']
)

add_workstream(
    'Workstream 20: Commission Visibility & Validation',
    'Andrea manages commissions in password-protected spreadsheets. Jodi needs to validate before release.',
    [
        'Build CRM-based commission view tied to closed-won deals and rep assignment',
        'Configure validation workflow: Jodi reviews and approves before release',
        'Create commission summary report by rep and time period',
    ],
    ['Visibility and validation only -- not a full commission calculation engine. Replaces manual spreadsheet review.']
)

add_workstream(
    'Workstream 21: NetSuite Integration Prep',
    'Jodi wants HubSpot data feeding NetSuite for month-end close. Clean CRM data is the prerequisite.',
    [
        'Document data readiness for API connection',
        'Create field mapping between HubSpot and NetSuite entities',
        'Identify data gaps that would block integration',
        'Produce integration readiness report with recommendations',
    ],
    ['PREP ONLY. Actual integration build requires separate SOW with NetSuite admin involvement. '
     'Does not include API development, middleware, or sync logic.']
)

add_workstream(
    'Workstream 22: Bill.com Integration Prep',
    'Bill.com handles AR/AP. Integration with HubSpot would automate revenue recognition.',
    [
        'Document AR/AP field mapping requirements',
        'Assess data readiness for future integration',
        'Produce integration readiness report',
    ],
    ['PREP ONLY. Actual build requires separate SOW.']
)

add_workstream(
    'Workstream 23: Seat/License Optimization',
    'Mixed versions, unused seats, and unclear allocation across hubs.',
    [
        'Audit all 23 current seats across core, sales, and service hubs',
        'Identify unused or misallocated seats',
        'Produce cost savings recommendation with specific seat changes',
        'Verify CPQ and sequence capabilities on current Enterprise tier',
    ],
    ['Deliverable: cost savings memo. Does not include license procurement.']
)

add_workstream(
    'Workstream 24: Advanced Reporting',
    'Phase 2 data (CPQ, marketing, enrichment) feeds into expanded reporting.',
    [
        'Configure automated board report delivery (scheduled email)',
        'Build pipeline coverage vs. quota tracking',
        'Add pace-to-goal reporting',
        'Create marketing campaign ROI reporting',
    ],
    ['Extends Phase 1 dashboards with Phase 2 data sources.']
)

# -- OPTIONAL ADD-ON --
doc.add_heading('Optional Add-On: Client Success Pipeline', level=2)
doc.add_paragraph(
    'A Client Success pipeline to track client health, renewals, and billing integration is available as an '
    'optional add-on. This workstream requires a separate discovery call with the Client Success team to '
    'define requirements and is not included in the base project cost.'
)
doc.add_paragraph('Estimated Add-On: 6-10 hours (median 8) / $900-$1,500 (median $1,200)')

# ============================================================
# 4. DELIVERABLES & TIMELINE
# ============================================================
doc.add_heading('4. Deliverables & Timeline', level=2)

doc.add_paragraph(
    'This engagement will be completed over a 10-14 week period using a phased approach. '
    'Phase 1 delivers the CRM foundation; Phase 2 builds CPQ, enrichment, and integrations on that foundation.'
)

add_table(
    ['Phase / Week Range', 'Scope Highlights'],
    [
        ('Phase 1: Weeks 1-2', 'CRM audit, contact classification rules, AI enrichment pilot, '
         'Outlook fix diagnosis, seat/license audit'),
        ('Phase 1: Weeks 3-4', 'Full contact classification and cleanup, data normalization, '
         'deduplication, duplicate protection'),
        ('Phase 1: Weeks 5-6', 'Pipeline restructuring, deal health scoring, automations, '
         'dashboards, Phase 1 UAT and go-live'),
        ('Phase 2: Weeks 7-9', 'CPQ product library build, quote templates, e-signatures, '
         'contract workflows, full database enrichment'),
        ('Phase 2: Weeks 10-12', 'Parent-child structure, email sequences, marketing automation, '
         'commission visibility, integration prep'),
        ('Phase 2: Weeks 13-14', 'Advanced reporting, Phase 2 training, UAT, go-live, '
         'documentation, hypercare begins'),
    ]
)

# ============================================================
# 5. ENGAGEMENT MODEL & PRICING
# ============================================================
doc.add_heading('5. Engagement Model & Pricing', level=2)

add_table(
    ['Component', 'Model', 'Fee'],
    [
        ('Phase 1: CRM Foundation', 'Fixed Fee', '$17,850'),
        ('Phase 2: CPQ + Enrichment + Integrations', 'Fixed Fee', '$14,550'),
        ('Project Total', 'Fixed Fee', '$32,400'),
    ]
)

doc.add_paragraph('Model: Fixed Fee')
doc.add_paragraph('Total Cost: $32,400')
doc.add_paragraph('Project Timeline: 10-14 weeks')

doc.add_paragraph('')
doc.add_paragraph('Payment Terms:')
doc.add_paragraph('Phase 1: 50% ($8,925) at Phase 1 kickoff, 50% ($8,925) at Phase 1 handoff')
doc.add_paragraph('Phase 2: 50% ($7,275) at Phase 2 kickoff, 50% ($7,275) at Phase 2 handoff')
doc.add_paragraph('Net-15 terms on all invoices')

doc.add_paragraph('')
doc.add_paragraph(
    'Technology & Administrative Fee: To help offset the costs of our technology stack and internal '
    'systems that improve project speed, quality, and efficiency, a 5% Technology & Administrative '
    'fee is applied to each invoice.'
)

doc.add_paragraph('')
doc.add_paragraph(
    'Budget Flexibility: Phase 1 can stand alone as a complete engagement. If budget is constrained, '
    'STRIVE can approve Phase 1 only and engage Phase 2 as a separate project after demonstrating '
    'Phase 1 ROI to the board.'
)

# ============================================================
# 6. PROJECT GOVERNANCE
# ============================================================
doc.add_heading('6. Project Governance', level=2)

doc.add_paragraph('Sayer Responsibilities')
for r in [
    'Lead configuration, QA, and implementation across all workstreams',
    'Conduct weekly check-ins and milestone reviews',
    'Provide documentation, training, and go-live support for each phase',
    'Manage project timeline and proactively communicate risks or blockers',
    'Deliver ROI tracking framework for post-project measurement',
    'Produce seat/license optimization memo with cost savings recommendations',
    'Deliver integration readiness reports for NetSuite and Bill.com',
]:
    doc.add_paragraph(r)

doc.add_paragraph('')
doc.add_paragraph('STRIVE Responsibilities')
for r in [
    'Designate 1 primary point of contact (Andrea recommended)',
    'Ensure Jodi and Andrea availability for discovery, UAT, and training',
    'Provide complete product/pricing inventory before Phase 2 CPQ build',
    'Provide contact classification rules and validate AI enrichment results',
    'Confirm deal stage definitions and probability framework',
    'Communicate standardized close won/lost reason categories',
    'Share 3-5 recent quotes/proposals across deal types for CPQ reference',
    'Mandate CRM usage across sales team (executive sponsorship critical)',
    'Complete UAT testing within agreed timeline for each phase',
    'Provide full HubSpot admin access',
    'Identify NetSuite admin contact for integration prep coordination',
    'Provide commission validation requirements and current spreadsheet structure',
]:
    doc.add_paragraph(r)

# ============================================================
# 7. ASSUMPTIONS & CONSTRAINTS
# ============================================================
doc.add_heading('7. Assumptions & Constraints', level=2)

for a in [
    'HubSpot Sales Hub Enterprise (existing instance -- reconfiguration)',
    'Up to 23 licensed users across core, sales, and service seats',
    'Outlook as primary email platform',
    'Up to 40 custom contact/company properties (expanded for CPQ and enrichment)',
    '1 active sales pipeline with pre-pipeline separated',
    '2 deal categories: partner-level and client-level',
    'AI enrichment using HubSpot Breeze credits (5-10K)',
    'Up to 12 automated workflows (expanded for CPQ and marketing)',
    '4-5 dashboards with up to 20 custom reports',
    'Complete product/pricing inventory provided before Phase 2',
    'Marketing Hub leveraged for campaigns (already licensed)',
    '10-14 week phased timeline',
    'Phase 1 can stand alone if Phase 2 is deferred',
    'All work conducted remotely',
    'Scope changes managed through formal change request',
]:
    doc.add_paragraph(a)

doc.add_paragraph('')
doc.add_paragraph('Out of Scope:')
for o in [
    'NetSuite integration build (prep only -- actual build requires separate SOW)',
    'Bill.com integration build (prep only)',
    'Client Success pipeline (available as optional add-on)',
    'Full commission calculation engine (visibility and validation only)',
    'Custom API development or middleware beyond HubSpot native',
    'Data migration from external systems (internal cleanup only)',
    'Ongoing managed services beyond 2-week hypercare per phase',
    'HubSpot license procurement or tier changes (recommendations provided)',
    'AI knowledge assistant for sales onboarding (future initiative)',
    'LinkedIn Sales Navigator integration (recommended for future)',
]:
    doc.add_paragraph(o)

# ============================================================
# 8. APPROVAL & NEXT STEPS
# ============================================================
doc.add_heading('8. Approval & Next Steps', level=2)

doc.add_paragraph(
    'Please confirm alignment on scope and structure so we can move forward with scheduling a '
    'formal kickoff and assigning your Sayer team.'
)
doc.add_paragraph(
    'If any scope areas, phasing, or pricing need to be adjusted, we welcome the conversation. '
    'This proposal is designed with budget flexibility -- Phase 1 can be approved independently '
    'with Phase 2 engaged after demonstrating ROI.'
)
doc.add_paragraph('')
doc.add_paragraph('To move forward with this engagement:')
doc.add_paragraph('Provide written approval of this proposal via email (full project or Phase 1 only)')
doc.add_paragraph('Schedule project kickoff within 1-2 weeks of written approval')
doc.add_paragraph('')
doc.add_paragraph(
    'Upon written approval, Sayer will issue a Master Services Agreement (MSA) via DocuSign, '
    'with this Scope of Work included. Once the MSA is executed, Sayer will finalize the '
    'implementation plan and begin onboarding.'
)

# Save
output_path = '/Users/harbuckconsulting/projects/Project Scoping Tool/Strive_Global/STRIVE Global Full CRM Overhaul Proposal.docx'
doc.save(output_path)
print(f'Option B proposal saved to: {output_path}')
