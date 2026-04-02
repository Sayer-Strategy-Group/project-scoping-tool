"""Generate STRIVE Global -- CRM Foundation Proposal (Option A)"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'


# ============================================================
# HELPERS
# ============================================================
def add_workstream(title, customer_story, approaches, assumptions):
    doc.add_heading(title, level=3)
    doc.add_paragraph('Customer Story')
    doc.add_paragraph(customer_story)
    doc.add_paragraph('Recommended Approach')
    for a in approaches:
        doc.add_paragraph(a)
    doc.add_paragraph('Assumptions')
    for a in assumptions:
        doc.add_paragraph(a)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph('')


# ============================================================
# TITLE
# ============================================================
doc.add_heading('STRIVE Global -- HubSpot CRM Foundation Proposal', level=1)
doc.add_paragraph('Option A: CRM Reconfiguration & Pipeline Optimization')

# ============================================================
# A NOTE ON THIS PROPOSAL
# ============================================================
doc.add_heading('A Note on This Proposal', level=3)
doc.add_paragraph(
    'Proposals are inherently a starting point -- not a final blueprint. This document reflects our '
    'understanding of your challenges, a proposed solution, and the resources required to execute successfully.'
)
doc.add_paragraph(
    'We recognize that you, as the client, live with these challenges daily, and there may be details '
    'we have yet to uncover. Our goal is to use this as a conversation starter -- to refine, align, and '
    'ensure we are solving the right problems in the right order.'
)
doc.add_paragraph(
    'If any aspect of the scope, solution, or pricing feels misaligned, let\'s use that as a basis '
    'for further discussion.'
)

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=2)

doc.add_paragraph(
    'STRIVE Global is a PE-backed (Essex Bay Capital) benefits technology company headquartered in Denver, CO '
    'with approximately 51 employees. STRIVE offers white-labeled employee benefits apps, AI-powered benefits '
    'assistants, scalable microsites, and a rewards platform (Divi) distributed through a broker channel model. '
    'Sales motions range from $6K client-level deals to $400K+ enterprise broker packs, creating significant '
    'pipeline variability that current CRM reporting cannot handle.'
)
doc.add_paragraph(
    'This proposal outlines a CRM Foundation engagement focused on reconfiguring STRIVE\'s existing HubSpot '
    'Sales Hub Enterprise instance to solve the most urgent operational pain points: separating pre-pipeline '
    'from active pipeline, classifying 10 years of uncoded contact data, implementing deal health scoring '
    'for objective forecasting, and delivering automated board-ready dashboards.'
)
doc.add_paragraph(
    'The engagement is structured as a 6-8 week implementation using a Foundation First approach: data cleanup '
    'and contact classification in weeks 1-4, followed by pipeline restructuring and dashboard delivery in weeks '
    '5-8. This sequencing ensures that dashboards are built on clean, trustworthy data from day one.'
)
doc.add_paragraph(
    'This scope is designed to stand alone as a complete engagement. It also serves as Phase 1 of a potential '
    'Full CRM Overhaul (Option B), which adds CPQ, marketing automation, and integration preparation.'
)

# ============================================================
# 2. OBJECTIVES
# ============================================================
doc.add_heading('2. Objectives', level=2)

objectives = [
    'Separate pre-pipeline deals from active pipeline to eliminate metric distortion (deal velocity, deal age, '
    'stage progression) and enable accurate forecasting.',
    'Classify all contact records by type (client, prospect, broker, partner) using AI enrichment to unlock '
    'segmented marketing, targeted outreach, and accurate reporting.',
    'Implement deal health scoring based on stage, time-in-stage, and activity recency to replace gut-feel '
    'forecasting with objective pipeline reviews.',
    'Deliver automated dashboards for partner-level, seller-level, and CEO/board-level reporting -- eliminating '
    'the 8-10 hours per month currently spent on manual report assembly.',
    'Standardize close won/lost reasons with categorized dropdowns to enable root cause analysis and coaching.',
    'Fix the Outlook email logging defect that causes emails to log across all deals associated with a contact.',
    'Establish data quality guardrails (deduplication, normalization, duplicate protection) to prevent data decay '
    'going forward.',
    'Create the operational foundation for a future Phase 2 engagement covering CPQ, marketing automation, '
    'and financial system integration preparation.',
]
for obj in objectives:
    doc.add_paragraph(obj)

# ============================================================
# 3. SCOPE OF WORK
# ============================================================
doc.add_heading('3. Scope of Work', level=2)
doc.add_paragraph('The following workstreams and tasks will be delivered during this engagement:')

# WORKSTREAM 1
add_workstream(
    'Workstream 1: CRM Architecture & Configuration Review',
    'STRIVE has an existing HubSpot Sales Hub Enterprise instance that has grown organically over several '
    'years. The configuration needs to be audited and optimized to support the team\'s actual workflows, '
    'with proper user roles, permissions, and team structure reflecting how the sales organization operates today.',
    [
        'Audit current HubSpot portal configuration, user roles, and permission sets',
        'Reconfigure team structure aligned to sales organization (partner vs. client sellers)',
        'Review and optimize seat allocation across sales, core, and service hubs',
        'Establish naming conventions and organizational standards for long-term CRM hygiene',
    ],
    [
        'Existing HubSpot Sales Hub Enterprise instance (reconfiguration, not net-new)',
        'Up to 23 licensed users across 3 core, 10 sales, and 10 service seats',
        'Portal configuration includes branding review, settings, and user provisioning',
    ]
)

# WORKSTREAM 2
add_workstream(
    'Workstream 2: Pipeline Restructuring',
    'STRIVE\'s current pipeline mixes pre-pipeline deals (FTA-0) with active deals, distorting key metrics '
    'like deal velocity and deal age. Deals that move backward from active pipeline to pre-pipeline further '
    'corrupt reporting. Andrea and Jodi need a clean separation between exploratory deals and committed pipeline, '
    'with partner-level and client-level segmentation for accurate reporting.',
    [
        'Separate pre-pipeline into its own pipeline or board view with dedicated stages',
        'Configure two deal categories within active pipeline: partner-level and client-level',
        'Map stage probabilities using Jodi\'s framework (5% pre-pipeline through 100% closed-won)',
        'Add Contracting stage between Commit (verbal) and Closed Won',
        'Configure required fields per stage to enforce data completeness',
        'Set up automated transitions from pre-pipeline to active pipeline (or closed-lost/revisit)',
        'Reset deal age when deal transitions from pre-pipeline to active pipeline',
    ],
    [
        'Two deal categories (partner, client) in same active pipeline -- same stages, different segmentation',
        'Pre-pipeline separated as distinct pipeline or board view',
        'Jodi\'s probability framework (5/20/40/50/60/80/90/100) as starting point, refined after first full cycle',
        'Product segmentation (microsite, AI, app, rewards) captured via product set fields, not pipeline structure',
    ]
)

# WORKSTREAM 3
add_workstream(
    'Workstream 3: Contact & Company Classification + Cleanup',
    'STRIVE has 10 years of contact records with zero classification. Every single record -- whether client, '
    'prospect, broker, or partner -- exists undifferentiated in the database. Andrea cannot run marketing '
    'campaigns, segment outreach, or report by contact type. This is the single largest data quality issue '
    'in the CRM and the foundation that every other improvement depends on.',
    [
        'Define contact type taxonomy: client, prospect, broker, partner (with sub-types as needed)',
        'Run AI enrichment pilot on 500 records using HubSpot Breeze credits to measure classification accuracy',
        'Build classification rules based on company associations, email domains, and deal history',
        'Execute full database classification based on pilot results',
        'Normalize data (capitalization, formatting, phone numbers, addresses)',
        'Deduplicate contacts using matching rules (email, name + company, phone)',
        'Fix contact-to-company associations where broken or missing',
        'Establish duplicate protection rules to prevent future data decay',
    ],
    [
        '10 years of uncoded data -- +30-50% data quality premium applied to estimates',
        'AI enrichment using HubSpot Breeze credits (5-10K credits available on Enterprise)',
        'Includes pilot enrichment on 500 records before full run',
        'Up to 30 custom contact and company properties',
        'Does not include external data source migration (internal cleanup only)',
    ]
)

# WORKSTREAM 4
add_workstream(
    'Workstream 4: Deal Health Scoring & Forecasting',
    'STRIVE has no algorithmic forecasting. Andrea and the team rely on gut feel and personal probability '
    'estimates. Jodi wants to replace this with an objective scoring system that tells the team -- and the '
    'board -- which deals are healthy, which are at risk, and which are stale. This also enables the 5x '
    'pipeline coverage and 20-25% close rate targets Jodi has set.',
    [
        'Create deal score formula based on: stage weight, time-in-stage, and activity recency',
        'Classify deals as Healthy, At Risk, or Stale based on scoring thresholds',
        'Replace free-form close won/lost reasons with standardized, categorized dropdowns',
        'Define granular close-lost reasons for root cause analysis (e.g., timing, budget, competitor, no decision)',
        'Configure pipeline views filtered by deal health status for objective pipeline reviews',
    ],
    [
        'Scoring formula and thresholds defined collaboratively with Andrea and Jodi',
        'Close won/lost categories defined before implementation',
        'Historical data used for calibration where reliable; industry benchmarks as fallback',
    ]
)

# WORKSTREAM 5
add_workstream(
    'Workstream 5: Reporting & Dashboards',
    'Andrea currently spends more than a full business day preparing each board report. Key metrics require '
    'manual calculations across multiple spreadsheets. Jodi needs a CEO dashboard that integrates sales, new '
    'business, and pipeline health. Andrea needs clonable dashboards she can filter by partner or seller for '
    '1:1 reviews.',
    [
        'Build partner-level dashboard with quick action filters (clonable per partner)',
        'Build seller-specific KPI dashboard (open deals, pipeline value, win rate, activity metrics)',
        'Build CEO/board dashboard (bookings, pipeline coverage, win rate, deal age, close analysis)',
        'Configure up to 12 custom reports across all dashboards',
        'Set up report scheduling and export capabilities for board deck preparation',
    ],
    [
        '3-4 dashboards: partner, seller, CEO/board, pipeline health',
        'Up to 12 custom reports across all dashboards',
        'Metrics include: open deals, pipeline value, avg deals/week, closed-won revenue, loss reasons, '
        'avg deal size, avg term length, appointments, contacts added, firm deals vs client deals breakdown',
        'Historical data available only from go-live date forward for activity-based metrics',
    ]
)

# WORKSTREAM 6
add_workstream(
    'Workstream 6: Automations & Workflows',
    'STRIVE currently has almost no sales automation beyond a broker trial form request. Andrea wants lead '
    'source tracking, microsite client capture, and stage-change notifications. The Outlook email logging '
    'defect -- where emails log to every deal a contact is associated with -- needs to be diagnosed and fixed.',
    [
        'Diagnose and fix Outlook email logging issue (emails logging to all associated deals)',
        'Configure lead source tracking automation (broker referral, inbound, self-prospected)',
        'Build form-to-CRM automation for microsite client capture (per-broker forms)',
        'Create deal stage-change notification workflows',
        'Set up duplicate protection rules (flash warning or block creation)',
        'Configure basic task reminder triggers for follow-up actions',
    ],
    [
        'Up to 8 automated workflows',
        'Outlook fix may require HubSpot support ticket if platform bug vs. configuration issue',
        'Lead source categories: broker referral, inbound, self-prospected',
        'Form automation depends on receiving microsite form link from Andrea',
    ]
)

# WORKSTREAM 7
add_workstream(
    'Workstream 7: Email & Communication Fix',
    'The Outlook-HubSpot integration creates significant noise by logging emails to every deal a contact is '
    'associated with. When a broker is connected to four deals, every email appears in all four deal activity '
    'feeds, making it impossible to understand the communication history for a specific opportunity.',
    [
        'Audit Outlook integration configuration across all users',
        'Configure email association rules to log at contact level with selective deal association',
        'Set up email domain exclusions for internal communications',
        'Validate fix with sample scenarios before rollout',
    ],
    [
        'Root cause diagnosis in week 1',
        'If configuration fix: 2-4 hours. If platform bug: HubSpot support escalation with unknown timeline',
        'Workaround documented if immediate fix not possible',
    ]
)

# WORKSTREAM 8
add_workstream(
    'Workstream 8: Training -- Admin',
    'Andrea and a designated admin need to be self-sufficient in managing the reconfigured CRM: running '
    'reports, managing pipeline hygiene, administering workflows, and maintaining data quality.',
    [
        'Deliver 1-2 admin training sessions covering pipeline management, reporting, workflow admin, and data hygiene',
        'Walk through AI enrichment tool usage and ongoing data quality monitoring',
        'Record sessions for future reference and new hire onboarding',
    ],
    [
        '1-2 sessions for Andrea + 1 designated admin',
        'Sessions delivered virtually via screen share',
        'Recordings provided for future onboarding use',
    ]
)

# WORKSTREAM 9
add_workstream(
    'Workstream 9: Training -- End Users',
    'The sales team needs to adopt the reconfigured CRM as their primary tool for deal tracking, activity '
    'logging, and pipeline management -- replacing the current pattern of prospecting in LinkedIn and logging '
    'activity after the fact.',
    [
        'Deliver 2-3 end-user training sessions covering daily CRM usage, deal entry, and activity logging',
        'Walk through dashboard consumption and self-service reporting',
        'Cover mobile app setup for on-the-go access',
        'Demonstrate how CRM data directly feeds their individual performance dashboards',
    ],
    [
        '2-3 sessions tailored to current team size',
        'Emphasis on showing reps how CRM data benefits them (not just management oversight)',
        'Mobile app walkthrough included',
    ]
)

# WORKSTREAM 10
add_workstream(
    'Workstream 10: Training -- Executive',
    'Jodi needs to consume board dashboards, pull scheduled reports, and understand forecast metrics without '
    'relying on Andrea to assemble reports manually.',
    [
        'Deliver 1 executive training session focused on CEO dashboard navigation',
        'Cover forecast consumption, pipeline health interpretation, and report scheduling',
        'Walk through export and formatting for board deck preparation',
    ],
    [
        '1 session focused on board reporting and forecast interpretation',
        'Includes report scheduling and automated delivery setup',
    ]
)

# WORKSTREAM 11
add_workstream(
    'Workstream 11: UAT, Go-Live & Documentation',
    'STRIVE needs a controlled go-live with testing and post-launch support to ensure adoption sticks. The '
    'documentation serves as the enablement foundation for future hires and ongoing CRM governance.',
    [
        'Develop test scripts for user acceptance testing across all configured workstreams',
        'Conduct 2 UAT sessions with Andrea, Jodi, and key stakeholders',
        'Execute bug fix window for issues identified during testing',
        'Create go-live cutover plan and checklist',
        'Deliver documentation: 1 admin guide, 1 end-user guide',
        'Provide 2-week hypercare support post-launch for issue resolution and adoption coaching',
    ],
    [
        'Includes test script creation, 2 UAT sessions, and bug fix window',
        'Go-live checklist and cutover plan included',
        'Deliverables: 1 admin guide, 1 end-user guide',
        '2 weeks post-launch hypercare support',
        'Ongoing managed services beyond hypercare not included',
    ]
)

# ============================================================
# 4. DELIVERABLES & TIMELINE
# ============================================================
doc.add_heading('4. Deliverables & Timeline', level=2)

doc.add_paragraph(
    'This engagement will be completed over a 6-8 week period using a Foundation First approach. '
    'During that time, Sayer will:'
)

deliverables = [
    'Audit and reconfigure HubSpot CRM with optimized user roles, permissions, and team structure',
    'Separate pre-pipeline from active pipeline with automated transitions',
    'Classify all contact records by type using AI enrichment and rule-based categorization',
    'Normalize, deduplicate, and establish ongoing data quality protections',
    'Implement deal health scoring and stage probability mapping',
    'Standardize close won/lost reasons with categorized dropdowns',
    'Fix Outlook email logging defect',
    'Deliver 3-4 automated dashboards (partner, seller, CEO/board, pipeline health)',
    'Build up to 8 automated workflows for lead tracking, notifications, and data quality',
    'Train admin, end-user, and executive teams with role-appropriate sessions',
    'Provide documentation, UAT, go-live support, and 2-week hypercare',
]
for d in deliverables:
    doc.add_paragraph(d)

add_table(
    ['Phase / Week Range', 'Scope Highlights'],
    [
        ('Weeks 1-2', 'CRM audit, contact classification rules, AI enrichment pilot (500 records), '
         'Outlook email fix diagnosis'),
        ('Weeks 3-4', 'Full contact classification and cleanup, data normalization, deduplication, '
         'duplicate protection rules'),
        ('Weeks 5-6', 'Pipeline restructuring, deal health scoring, automations, dashboard build'),
        ('Weeks 7-8', 'Reporting refinement, training sessions (admin, end-user, executive), '
         'UAT, go-live, documentation, hypercare begins'),
    ]
)

# ============================================================
# 5. ENGAGEMENT MODEL & PRICING
# ============================================================
doc.add_heading('5. Engagement Model & Pricing', level=2)

add_table(
    ['Component', 'Model', 'Fee'],
    [('Project Delivery', 'Fixed Fee', '$17,850')]
)

doc.add_paragraph('Model: Fixed Fee')
doc.add_paragraph('Total Cost: $17,850')
doc.add_paragraph('Project Timeline: 6-8 weeks')

doc.add_paragraph('')
doc.add_paragraph('Payment Terms:')
doc.add_paragraph('50% ($8,925) invoiced at project start')
doc.add_paragraph('50% ($8,925) invoiced at final handoff (~Week 6-8)')
doc.add_paragraph('Net-15 terms on all invoices')

doc.add_paragraph('')
doc.add_paragraph(
    'Technology & Administrative Fee: To help offset the costs of our technology stack and internal '
    'systems that improve project speed, quality, and efficiency, a 5% Technology & Administrative '
    'fee is applied to each invoice.'
)

# ============================================================
# 6. PROJECT GOVERNANCE
# ============================================================
doc.add_heading('6. Project Governance', level=2)

doc.add_paragraph('Sayer Responsibilities')
for r in [
    'Lead configuration, QA, and implementation across all workstreams',
    'Conduct weekly check-ins and milestone reviews',
    'Provide documentation, training, and go-live support',
    'Manage project timeline and proactively communicate risks or blockers',
    'Deliver ROI tracking framework for post-project measurement',
]:
    doc.add_paragraph(r)

doc.add_paragraph('')
doc.add_paragraph('STRIVE Responsibilities')
for r in [
    'Designate 1 primary point of contact for decisions and approvals (Andrea recommended)',
    'Ensure Jodi and Andrea availability for discovery, UAT, and training sessions',
    'Provide contact classification rules and validate AI enrichment results',
    'Confirm deal stage definitions and probability framework before pipeline build',
    'Communicate standardized close won/lost reason categories',
    'Mandate CRM usage across sales team -- executive sponsorship from Jodi is critical for adoption',
    'Complete UAT testing within agreed timeline',
    'Provide full HubSpot admin access',
]:
    doc.add_paragraph(r)

# ============================================================
# 7. ASSUMPTIONS & CONSTRAINTS
# ============================================================
doc.add_heading('7. Assumptions & Constraints', level=2)

for a in [
    'HubSpot Sales Hub Enterprise (existing instance -- reconfiguration, not net-new setup)',
    'Up to 23 licensed users across 3 core, 10 sales, and 10 service seats',
    'Outlook as primary email platform',
    'Up to 30 custom contact/company properties',
    '1 active sales pipeline with pre-pipeline separated',
    '2 deal categories: partner-level and client-level (same stages, different segmentation)',
    'Stage probability mapping using Jodi\'s framework as starting point',
    'AI enrichment using HubSpot Breeze credits (5-10K credits available)',
    'Up to 8 automated workflows',
    '3-4 dashboards with up to 12 custom reports',
    '6-8 week implementation timeline',
    'All work conducted remotely unless otherwise agreed',
    'Any scope changes managed through formal change request and may impact cost or timing',
]:
    doc.add_paragraph(a)

doc.add_paragraph('')
doc.add_paragraph('Out of Scope:')
for o in [
    'CPQ / quoting / product library (see Option B: Full CRM Overhaul)',
    'Email sequencing or sales cadence automation',
    'Marketing automation (newsletters, campaigns, landing pages)',
    'NetSuite, Bill.com, or any ERP/financial system integration',
    'Client Success pipeline or client health scoring',
    'Commission tracking or calculation',
    'Contract management or e-signature workflows',
    'Custom API development or middleware (n8n/Zapier)',
    'Data migration from external systems (HubSpot internal cleanup only)',
    'Ongoing managed services beyond 2-week hypercare',
    'HubSpot license procurement or tier changes',
]:
    doc.add_paragraph(o)

# ============================================================
# 8. APPROVAL & NEXT STEPS
# ============================================================
doc.add_heading('8. Approval & Next Steps', level=2)

doc.add_paragraph(
    'Please confirm alignment on scope and structure so we can move forward with scheduling a '
    'formal kickoff and assigning your Sayer team.'
)
doc.add_paragraph(
    'If any scope areas, phasing, or pricing need to be adjusted to fit your needs, we welcome '
    'the conversation.'
)
doc.add_paragraph('')
doc.add_paragraph('To move forward with this engagement:')
doc.add_paragraph('Provide written approval of this proposal via email')
doc.add_paragraph('Schedule project kickoff within 1-2 weeks of written approval')
doc.add_paragraph('')
doc.add_paragraph(
    'Upon written approval, Sayer will issue a Master Services Agreement (MSA) via DocuSign, '
    'with this Scope of Work included. Once the MSA is executed, Sayer will finalize the '
    'implementation plan and begin onboarding.'
)

# Save
output_path = '/Users/harbuckconsulting/projects/Project Scoping Tool/Strive_Global/STRIVE Global CRM Foundation Proposal.docx'
doc.save(output_path)
print(f'Option A proposal saved to: {output_path}')
